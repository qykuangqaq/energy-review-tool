# -*- coding: utf-8 -*-
import sys
import os
import io

# ===================== 编码设置 =====================
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
if sys.stderr.encoding != 'utf-8':
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

os.environ['PYTHONIOENCODING'] = 'utf-8'
os.environ['PYTHONUTF8'] = '1'

import logging
logging.getLogger("openai").setLevel(logging.ERROR)
logging.getLogger("httpx").setLevel(logging.ERROR)
logging.getLogger("httpcore").setLevel(logging.ERROR)

import streamlit as st
from openai import OpenAI
import docx
import pdfplumber

# ===================== 页面配置 =====================
st.set_page_config(page_title="节能报告初审工具", layout="wide")

# ===================== 初始化 session_state =====================
if "review_result" not in st.session_state:
    st.session_state.review_result = None
if "file_content" not in st.session_state:
    st.session_state.file_content = None
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# ===================== 访问密码校验 =====================
def check_password():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False

    if st.session_state.authenticated:
        return True

    st.title("🔐 请输入访问密码")
    password = st.text_input("密码", type="password", key="login_password")
    if st.button("验证"):
        if password == st.secrets["APP_PASSWORD"]:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("密码错误，请重试")
    return False

if not check_password():
    st.stop()

# ===================== 主界面 =====================
st.title("📋 新上项目（在建整改项目）节能报告初审工具")
st.markdown("上传节能报告（**仅支持 .pdf 或 .docx**，旧版 .doc 请先另存为 .docx）")

# ===================== 系统提示词（完整14项评审规则） =====================
system_prompt = """
你是一位专业的节能评审专家。请严格按照下面的评审表，对用户提供的节能报告逐项进行审核，并输出评审结论和具体的修改建议。

## 评审项目及要求

### 1. 章节完整性（否决项，缺少任何一项直接不通过）
必须包含以下全部章节：
- 一、项目摘要表
- 二、项目基本情况
- 三、分析评价依据
- 四、建设方案节能分析和比选
- 五、节能降碳措施
- 六、能源消费情况核算及能效水平评价（含碳排放评价）
- 七、能源消费影响分析
- 八、结论
- 九、节能承诺书

### 2. 项目边界一致性
- 报告描述（项目名称/建设单位/建设内容/面积）是否与备案证或环评批复等文件一致；
- 偏差超过 20% 需要更换备案证或补充详细说明；
- 评估是否有拆分项目或合并项目的嫌疑。

### 3. 项目产业政策相符性
- 对照《广东省“两高”项目管理目录（2025年版）》判断是否属于“两高”；
- 对照《产业结构调整指导目录》《广东省大气污染防治条例》《市场准入负面清单（2022年版）》《鼓励外商投资产业目录（2022年版）》核实工艺/技术是否属限制类或淘汰类。

### 4. 主要用能设备能效水平
- 主要用能设备是否属《产业结构调整指导目录》限制类/淘汰类；
- 是否采用了《高耗能落后机电设备（产品）淘汰目录》（第一至四批）中的设备；
- 能效水平是否满足限额标准准入等级或《重点用能产品设备能效先进水平、节能水平和准入水平（2022年版）》的节能水平。

### 5. 主要能效指标先进性
- 单耗原则上需满足现行标准的先进值；如无现行标准，是否与行业先进值对比并说明先进值来源。

### 6. 建设方案先进性（完善项，不直接导致不通过）
- 项目建设方案描述是否合理可行；
- 工艺技术方案是否有节能分析比选；
- 节能先进性特征是否简洁明了，是否有量化对比指标。

### 7. 总平面布置
- 平面布置图是否清晰；
- 平面布置方案是否结合项目实际情况进行具体分析。

### 8. 工艺技术先进性
- 主要用能工艺/工序识别是否到位合理；
- 是否对主要用能工艺/工序做了量化节能比选分析；特定行业是否对照国家/行业/地方限额/定额标准进行对标。

### 9. 用能设备配备合理性
- 结合产品方案、产能规模、工艺路径，评估设备配备数量/装机容量/备用情况是否合理。

### 10. 节能措施可实施性
- 是否提出了具体的、有针对性的节能技术措施和管理措施；
- 节能措施是否有量化的节能效果；
- 是否结合当前政策积极开展可再生能源开发利用。

### 11. 能耗计算合理性
- 项目边界内的能源消费需求是否均已统计；
- 生产、照明、供配电、给排水、插座、办公等系统能耗计算是否准确、合理、可行；
- 是否开展了能量平衡/热量平衡计算，能源需求是否有支撑；
- 折标系数取值是否合理，特别是煤/油/气等需实测热值的能源是否结合实际情况评估。

### 12. 能效指标先进性
- 能效指标计算过程是否详尽、可溯源、可考究；
- 对标对象/限额标准是否有详尽描述，来源是否充分；
- 对标过程是否合理，结论是否满足新上项目要求；
- 原则上需满足先进值（1级）能效要求。

### 13. 碳排放评价
- 碳排放评价依据出处是否清晰、准确、统一；
- 碳排放量核算方法是否合理，计算过程是否完整；
- 碳排放强度指标计算是否正确，是否有对标过程。

### 14. 对区域能耗“双控”的影响
- 区域基准期能源消费总量/增量和能耗强度指标是否准确；
- 区域预测期能源消费总量/增量和能耗强度下降指标是否准确；
- 项目 m、n 值计算过程是否详尽，计算结论是否合理。

## 输出格式（必须严格遵守）
请严格按以下结构输出结果，不得遗漏任何一项：

**初审结论**：【通过 / 基本通过 / 不通过】

**逐项评审意见**：
1. 章节完整性：
2. 项目边界一致性：
3. 项目产业政策相符性：
4. 主要用能设备能效水平：
5. 主要能效指标先进性：
6. 建设方案先进性（完善项）：
7. 总平面布置：
8. 工艺技术先进性：
9. 用能设备配备合理性：
10. 节能措施可实施性：
11. 能耗计算合理性：
12. 能效指标先进性：
13. 碳排放评价：
14. 对区域能耗“双控”的影响：

**综合建议**：（汇总改进方向；如存在否决项或严重问题，必须明确指出）
"""

# ===================== 文件上传 =====================
uploaded_file = st.file_uploader(
    "上传节能报告（支持 .pdf / .docx）",
    type=["pdf", "docx"],
    key="file_uploader"
)

# ===================== 额外指令词输入框 =====================
st.markdown("---")
st.markdown("### ⚙️ 针对本次评审的特殊要求（可选）")
extra_instructions = st.text_area(
    "填写后，会覆盖或补充默认评审规则。例如：\"本次不审查碳排放部分\"、\"重点审查用能设备能效水平\"、\"只关注第3、5、11项\"等。",
    placeholder="请输入额外的评审要求（留空则按默认规则评审）",
    height=100
)

# ===================== 评审按钮与逻辑 =====================
if uploaded_file is not None:
    st.success(f"已上传文件：{uploaded_file.name}（大小：{uploaded_file.size / 1024:.1f} KB）")

    start_review = st.button("🚀 开始评审", type="primary")

    if start_review:
        # 清空旧的聊天记录（新评审开始）
        st.session_state.chat_history = []

        st.write("📖 正在读取文件内容……")
        file_content = ""
        try:
            if uploaded_file.name.endswith(".pdf"):
                with pdfplumber.open(uploaded_file) as pdf:
                    total_pages = len(pdf.pages)
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    for i, page in enumerate(pdf.pages):
                        status_text.text(f"解析 PDF 第 {i+1}/{total_pages} 页……")
                        page_text = page.extract_text()
                        if page_text:
                            file_content += page_text + "\n"
                        progress_bar.progress((i + 1) / total_pages)
                    status_text.text("PDF 解析完成")
            else:  # docx
                doc = docx.Document(uploaded_file)
                paragraphs = doc.paragraphs
                total_paras = len(paragraphs)
                progress_bar = st.progress(0)
                status_text = st.empty()
                for i, para in enumerate(paragraphs):
                    status_text.text(f"解析 Word 段落 {i+1}/{total_paras}……")
                    file_content += para.text + "\n"
                    progress_bar.progress((i + 1) / total_paras)
                status_text.text("Word 解析完成")
        except Exception as e:
            st.error(f"文件读取失败：{e}")
            st.stop()

        if not file_content.strip():
            st.warning("⚠️ 未能提取到文字内容，请确认文件不是扫描图片且不是空白文件。")
        else:
            st.success(f"文件读取成功，共提取 {len(file_content)} 个字符。正在调用 AI 评审……")

            # --- 组合最终的系统提示词 ---
            final_prompt = system_prompt
            if extra_instructions.strip():
                final_prompt += f"\n\n## ⚠️ 本次评审特别要求\n{extra_instructions.strip()}\n请严格遵循以上特别要求，其优先级高于默认规则。"

            client = OpenAI(
                api_key=st.secrets["API_KEY"],
                base_url="https://api.deepseek.com"
            )
            try:
                with st.spinner("AI 评审中，请耐心等待（约 30-90 秒）……"):
                    response = client.chat.completions.create(
                        model="deepseek-v4-pro",
                        messages=[
                            {"role": "system", "content": final_prompt},
                            {"role": "user", "content": file_content}
                        ],
                        temperature=0.2,
                        max_tokens=4096,
                    )
                result = response.choices[0].message.content

                # 将结果和文件内容存入 session_state，防止页面刷新丢失
                st.session_state.review_result = result
                st.session_state.file_content = file_content
                st.session_state.reviewed_file_name = uploaded_file.name

                st.success("✅ 评审完成！")
                st.markdown(result)

                # ==================== 生成 Word 下载按钮 ====================
                from docx import Document

                word_doc = Document()
                word_doc.add_heading('节能报告初审意见', level=1)

                for line in result.split('\n'):
                    if line.strip():
                        word_doc.add_paragraph(line.strip())

                word_buffer = io.BytesIO()
                word_doc.save(word_buffer)
                word_buffer.seek(0)

                st.download_button(
                    label="📥 下载评审意见 (Word)",
                    data=word_buffer,
                    file_name=f"评审意见_{uploaded_file.name.rsplit('.', 1)[0]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"调用 API 出错：{e}")

# ===================== 如果 session 中已有评审结果，直接显示（防清空） =====================
elif st.session_state.review_result is not None:
    st.success("✅ 评审完成（历史结果）")
    st.markdown(st.session_state.review_result)

    # 重新生成下载按钮（避免因刷新丢失）
    from docx import Document
    word_doc = Document()
    word_doc.add_heading('节能报告初审意见', level=1)
    for line in st.session_state.review_result.split('\n'):
        if line.strip():
            word_doc.add_paragraph(line.strip())
    word_buffer = io.BytesIO()
    word_doc.save(word_buffer)
    word_buffer.seek(0)
    st.download_button(
        label="📥 下载评审意见 (Word)",
        data=word_buffer,
        file_name=f"评审意见_{st.session_state.reviewed_file_name.rsplit('.', 1)[0]}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ===================== 侧边栏：问答功能 =====================
st.sidebar.title("💬 对报告提问")
if st.session_state.review_result is None:
    st.sidebar.info("请先完成报告评审，然后可以在此处针对报告内容提问。")
else:
    # 显示聊天记录
    for msg in st.session_state.chat_history:
        if msg["role"] == "user":
            st.sidebar.markdown(f"**🧑 你：** {msg['content']}")
        else:
            st.sidebar.markdown(f"**🤖 AI：** {msg['content']}")

    # 输入新问题
    with st.sidebar.form("question_form", clear_on_submit=True):
        user_question = st.text_input("输入你的问题", key="question_input")
        submit_question = st.form_submit_button("发送")

    if submit_question and user_question.strip():
        # 将用户问题加入历史
        st.session_state.chat_history.append({"role": "user", "content": user_question.strip()})

        # 构建问答用的系统提示词（包含报告背景和评审结果）
        qa_system_prompt = f"""你是一位精通节能评审的专家。以下是一份节能报告的评审结果，用户将基于这份报告和评审意见提出问题。请结合报告内容和评审结果，给出专业、准确的回答。

## 报告内容
{st.session_state.file_content}

## 评审结果
{st.session_state.review_result}

请根据以上信息回答用户的问题。如果问题超出报告范围，请如实说明。"""

        # 调用 API
        client = OpenAI(
            api_key=st.secrets["API_KEY"],
            base_url="https://api.deepseek.com"
        )
        with st.sidebar.spinner("思考中…"):
            try:
                response = client.chat.completions.create(
                    model="deepseek-v4-pro",
                    messages=[
                        {"role": "system", "content": qa_system_prompt},
                        *st.session_state.chat_history
                    ],
                    temperature=0.2,
                    max_tokens=2048,
                )
                answer = response.choices[0].message.content
                st.session_state.chat_history.append({"role": "assistant", "content": answer})
                st.rerun()  # 刷新以显示新消息
            except Exception as e:
                st.sidebar.error(f"提问失败：{e}")